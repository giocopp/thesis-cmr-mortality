"""
Generate Word document summarizing Model A analysis results.
Output: documents/model_a_progress_report.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ── Title ───────────────────────────────────────────────────────────
title = doc.add_heading("Model A Progress Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Weather x Post-MoU Interaction Analysis: Results, Diagnostics, and Interpretation"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Giorgio Coppola | MDS Thesis, Hertie School | March 2025"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── 1. Research Design ──────────────────────────────────────────────
doc.add_heading("1. Research Design", level=1)

doc.add_paragraph(
    "The thesis asks whether the 2017 EU-Libya Memorandum of Understanding (MoU) "
    "made the Central Mediterranean Route (CMR) more dangerous for migrants. "
    "The primary analysis (Model A) tests a specific mechanism: did the MoU "
    "amplify the lethality of adverse sea conditions?"
)

doc.add_heading("1.1 Estimand", level=2)
doc.add_paragraph(
    "The theoretical estimand is the change in the weather-mortality gradient: "
    "how much more sensitive did per-incident mortality become to weather "
    "conditions after the MoU? This is a causal interaction, not a standard ATT. "
    "On the multiplicative scale, exp(beta_3) is the ratio of the post-MoU "
    "weather IRR to the pre-MoU weather IRR."
)

doc.add_heading("1.2 Model Specification", level=2)
doc.add_paragraph(
    "dead_missing_i ~ NegBin(mu_i), "
    "log(mu_i) = alpha_g + lambda_t + beta_3(Weather_i x Post_t) + X'gamma"
)
doc.add_paragraph(
    "where dead_missing is the total dead + missing per incident (>= 1 for all "
    "CMR incidents by construction), alpha_g is a grid fixed effect (1-degree), "
    "lambda_t includes year and month-of-year FE, Weather is SWH, wind speed, "
    "or wind gust at the incident location on the day of the incident (ERA5), "
    "and Post is a binary indicator for dates >= Feb 2, 2017."
)

doc.add_heading("1.3 Identification Assumptions", level=2)
doc.add_paragraph(
    "1. Weather exogeneity conditional on FE (physics determines weather, not policy)\n"
    "2. MoU exogeneity (no parents in the DAG)\n"
    "3. Gradient stability: absent the MoU, the weather-mortality gradient "
    "would have remained stable over time (testable with placebos)\n"
    "4. Exclusion restriction for the interaction\n"
    "5. SUTVA\n"
    "6. No anticipation"
)

doc.add_heading("1.4 Why No Control Routes", level=2)
doc.add_paragraph(
    "The Eastern Mediterranean Route is contaminated by the 2016 EU-Turkey deal. "
    "The Western Mediterranean and West African Atlantic routes have fundamentally "
    "different climate, sea conditions, and smuggler networks. Cross-route parallel "
    "gradient trends are not credible. The analysis relies on within-CMR temporal "
    "variation and placebo treatment dates."
)

# ── 2. Data ─────────────────────────────────────────────────────────
doc.add_heading("2. Data", level=1)

doc.add_heading("2.1 IOM Missing Migrants Project", level=2)
doc.add_paragraph(
    "1,381 CMR incidents (type 'Incident' only), geocoded, 2014-2025. "
    "Outcome: dead + missing per incident (min = 1, median = 2, mean = 14.4, max = 550). "
    "Var/Mean ratio = 113.9 (strong overdispersion). "
    "Pre-MoU: 239 incidents. Post-MoU: 1,142 incidents."
)

doc.add_heading("2.2 ERA5 Weather", level=2)
doc.add_paragraph(
    "Daily weather extracted at each incident's nearest grid cell on the day of "
    "the incident (day 0) and 7 preceding days. Variables: significant wave height "
    "(SWH, 0.5 deg), wind speed (u10/v10 -> speed, 0.25 deg), 10-m wind gust "
    "(i10fg, 0.25 deg), SST (0.25 deg), mean wave period (0.5 deg), total "
    "precipitation (0.25 deg). Max wave height (hmax) had a grid index bug "
    "(now fixed, pending re-extraction)."
)

doc.add_heading("2.3 Estimation Sample", level=2)
doc.add_paragraph(
    "After dropping 149 incidents with missing SWH (coastal grid cells without wave "
    "data) and 25 singleton grid cells dropped by fixest: N = 1,207 incidents "
    "(211 pre-MoU, 996 post-MoU) for the primary specification."
)

# ── 3. Weather Overlap Diagnostic ──────────────────────────────────
doc.add_heading("3. Weather Overlap Diagnostic", level=1)

doc.add_paragraph(
    "A prerequisite for the interaction design: do incidents occur under comparable "
    "weather conditions in both periods? If post-MoU incidents only occur in calm "
    "seas, beta_3 is identified off extrapolation."
)

headers = ["Variable", "n pre", "Mean pre", "Median pre",
           "n post", "Mean post", "Median post", "KS stat", "KS p"]
rows = [
    ["SWH (m)", "211", "0.81", "0.67", "1021", "0.69", "0.55", "0.169", "<0.001"],
    ["Wind (m/s)", "239", "4.84", "4.64", "1142", "4.69", "4.42", "0.062", "0.427"],
    ["Gust (m/s)", "239", "7.31", "6.85", "1142", "7.61", "7.22", "0.071", "0.278"],
    ["Wave period (s)", "211", "4.66", "4.40", "1021", "4.29", "4.11", "0.178", "<0.001"],
    ["SST (C)", "188", "22.2", "23.0", "634", "21.8", "21.6", "0.097", "0.130"],
]
add_table(doc, headers, rows)

doc.add_paragraph()
doc.add_paragraph(
    "Result: Common support holds. For all variables, the pre-MoU range sits "
    "within the post-MoU range. Wind speed and gust show no distributional shift. "
    "SWH and wave period shift modestly toward calmer seas post-MoU (consistent "
    "with literature on boat degradation). The interaction is identified over "
    "shared support, not extrapolation."
)

# Add figure
fig1_path = os.path.join(BASE_DIR, "output", "figures", "weather_overlap_densities.png")
if os.path.exists(fig1_path):
    doc.add_paragraph()
    p_fig = doc.add_paragraph()
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fig.add_run()
    run.add_picture(fig1_path, width=Inches(6.0))
    cap = doc.add_paragraph("Figure 1: Weather conditions at incident locations, pre vs post MoU.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)

# ── 4. Model A Results ─────────────────────────────────────────────
doc.add_heading("4. Model A Results", level=1)

doc.add_paragraph(
    "All models estimated with fixest::fenegbin (proper high-dimensional FE handling, "
    "no convergence issues). Heteroskedasticity-robust SEs unless noted. "
    "1-degree grid FE + year FE + month-of-year FE."
)

doc.add_heading("4.1 Primary Specifications", level=2)

headers = ["Specification", "beta_3", "SE", "p-value", "exp(beta_3)"]
rows = [
    ["A: SWH x Post", "-0.204", "0.204", "0.319", "0.816"],
    ["B: Wind x Post", "-0.053", "0.051", "0.294", "0.948"],
    ["C: Joint (SWH + Wind)", "-0.101 / -0.035", "0.333 / 0.083", "0.762 / 0.668", "0.904 / 0.965"],
    ["D: Gust x Post", "-0.092", "0.038", "0.016 **", "0.913"],
]
add_table(doc, headers, rows)

doc.add_paragraph()
doc.add_paragraph(
    "The SWH and wind interactions are not significant. The gust interaction "
    "is significant at 5% (p = 0.016): a 1 m/s increase in gust intensity is "
    "associated with 8.7% fewer deaths per incident post-MoU compared to pre-MoU. "
    "All signs are negative -- the gradient decreased, opposite to the hypothesis."
)

doc.add_heading("4.2 Sensitivity and Robustness", level=2)

headers = ["Specification", "beta_3", "SE", "p-value", "exp(beta_3)"]
rows = [
    ["A: no grid FE", "-0.301", "0.230", "0.191", "0.740"],
    ["A: 0.25-deg grid FE", "-0.104", "0.244", "0.668", "0.901"],
    ["A: month-only FE", "-0.044", "0.227", "0.845", "0.957"],
    ["A: trim <= 100", "-0.125", "0.199", "0.531", "0.883"],
    ["D: trim <= 100", "-0.056", "0.036", "0.124", "0.946"],
    ["A: cluster SE (grid)", "-0.204", "0.194", "0.294", "0.816"],
    ["D: cluster SE (grid)", "-0.092", "0.044", "0.039 **", "0.913"],
    ["A: Poisson", "-0.331", "0.240", "0.168", "0.718"],
    ["D: Poisson", "-0.076", "0.048", "0.116", "0.927"],
]
add_table(doc, headers, rows)

doc.add_paragraph()

doc.add_paragraph(
    "Key findings from sensitivity analysis:"
)

points = [
    "The SWH interaction is robustly null across all specifications "
    "(grid resolution, temporal FE, outlier trimming, SE clustering, Poisson).",
    "The gust interaction survives grid-clustered SEs (p = 0.039) but does NOT "
    "survive outlier trimming (p = 0.124) or Poisson estimation (p = 0.116). "
    "It is not robust.",
    "Outlier trimming (dead+missing > 100) drops 33 incidents, all pre-MoU mass-casualty "
    "events. The top 5 deadliest incidents (400-550 dead+missing) are all pre-MoU.",
]
for pt in points:
    doc.add_paragraph(pt, style="List Bullet")

# ── 5. Placebo Tests ───────────────────────────────────────────────
doc.add_heading("5. Placebo and Stability Tests", level=1)

doc.add_heading("5.1 Placebo Treatment Dates -- SWH", level=2)

doc.add_paragraph(
    "The SWH x Post interaction was re-estimated at 14 placebo treatment dates "
    "(quarterly, 2015-Q1 to 2018-Q3). If the MoU date is special, beta_3 should "
    "stand out against placebo estimates."
)

headers = ["Placebo date", "beta_3", "SE", "p-value", "n pre", "n post"]
rows = [
    ["2015-01", "-0.312", "1.499", "0.835", "43", "1189"],
    ["2015-04", "-0.850", "0.287", "0.003 **", "49", "1183"],
    ["2015-07", "-0.407", "0.272", "0.135", "64", "1168"],
    ["2016-01", "-0.302", "0.253", "0.233", "88", "1144"],
    ["2016-07", "-0.069", "0.274", "0.801", "116", "1116"],
    ["2017-02 (MoU)", "-0.204", "0.204", "0.319", "211", "1021"],
    ["2017-07", "-0.173", "0.204", "0.395", "314", "918"],
    ["2018-01", "-0.147", "0.205", "0.475", "346", "886"],
    ["2018-07", "-0.120", "0.202", "0.551", "373", "859"],
]
add_table(doc, headers, rows)

doc.add_paragraph()
doc.add_paragraph(
    "Result: The MoU date is not special. beta_3 hovers between -0.04 and -0.30 "
    "at all dates from 2016 onward, always with p > 0.19. The coefficient path "
    "is flat -- there is no break at February 2017. The one significant placebo "
    "(April 2015, p = 0.003) is driven by only 49 pre-period observations and "
    "a convergence warning."
)

doc.add_heading("5.2 Placebo Treatment Dates -- Gust", level=2)

headers = ["Placebo date", "beta_3", "SE", "p-value"]
rows = [
    ["2015-04", "-0.095", "0.099", "0.337"],
    ["2016-01", "-0.076", "0.060", "0.202"],
    ["2016-07", "-0.076", "0.055", "0.164"],
    ["2016-10", "-0.085", "0.052", "0.107"],
    ["2017-02 (MoU)", "-0.092", "0.038", "0.016 **"],
    ["2017-07", "-0.084", "0.037", "0.024 **"],
    ["2017-10", "-0.072", "0.036", "0.047 **"],
    ["2018-04", "-0.088", "0.035", "0.013 **"],
    ["2018-07", "-0.071", "0.035", "0.043 **"],
]
add_table(doc, headers, rows)

doc.add_paragraph()
doc.add_paragraph(
    "Result: The gust interaction is significant at the MoU date (p = 0.016) but "
    "also at every date from mid-2017 through 2018. The coefficient is approximately "
    "constant at -0.07 to -0.09 from 2015 onward. This is a gradual pre-existing "
    "trend, not a policy break. The MoU date does not stand out."
)

# Add figure
fig2_path = os.path.join(BASE_DIR, "output", "figures", "placebo_beta3.png")
if os.path.exists(fig2_path):
    doc.add_paragraph()
    p_fig = doc.add_paragraph()
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fig.add_run()
    run.add_picture(fig2_path, width=Inches(5.5))
    cap = doc.add_paragraph("Figure 2: Placebo treatment dates. Red dashed line = actual MoU (Feb 2017). Diamond = true estimate.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)

doc.add_heading("5.3 Pre-Period Gradient Stability", level=2)

doc.add_paragraph(
    "The SWH-mortality gradient was estimated separately for each pre-MoU year "
    "(Poisson, no grid FE due to small samples):"
)

headers = ["Year", "n", "SWH slope", "exp(slope)"]
rows = [
    ["2014", "43", "-0.219", "0.804"],
    ["2015", "45", "+0.964", "2.622"],
    ["2016", "110", "-0.274", "0.760"],
    ["2017 (Jan)", "13", "-0.889", "0.411"],
]
add_table(doc, headers, rows)

doc.add_paragraph()
doc.add_paragraph(
    "The gradient is not stable: 2015 shows a strongly positive slope (higher "
    "waves associated with more deaths), while 2014 and 2016 show negative slopes. "
    "A Wald test for joint significance of year interactions rejects at the 10% "
    "level (p = 0.083). This weakens the gradient stability assumption (Assumption 3)."
)

doc.add_paragraph(
    "Caveat: these are Poisson estimates with very small samples (43-110 obs) "
    "and no spatial FE. The instability may reflect noise from small samples "
    "rather than genuine gradient variation. However, the evidence does not "
    "support gradient stability."
)

# ── 6. Interpretation ──────────────────────────────────────────────
doc.add_heading("6. Interpretation", level=1)

doc.add_heading("6.1 What the Results Say", level=2)

doc.add_paragraph(
    "The weather x post-MoU interaction design does not detect evidence that "
    "the MoU changed the weather-mortality gradient on the CMR. The primary "
    "SWH interaction is null (beta_3 = -0.20, p = 0.32), robust across all "
    "specifications. The gust interaction, while nominally significant, fails "
    "the placebo test (it is a pre-existing trend) and does not survive outlier "
    "trimming."
)

doc.add_heading("6.2 Possible Explanations", level=2)

points = [
    ("The MoU changed the type of danger, not the weather sensitivity. "
     "Pre-MoU: large boats, open-sea crossings, weather-driven mass-casualty "
     "shipwrecks (the top 5 deadliest events are all pre-MoU), SAR present. "
     "Post-MoU: small fragile boats, short-range crossings near Libya in calm "
     "seas, deaths from overcrowding, engine failure, dehydration, and LCG "
     "interceptions. The weather gradient flattened because weather is no "
     "longer the dominant killer."),
    ("Detection bias. Post-MoU, the worst weather incidents (sinkings far "
     "from shore, no witnesses) may go undetected. If these 'ghost boats' "
     "have high mortality in bad weather, their absence from the data would "
     "attenuate beta_3 toward zero or even reverse it."),
    ("Power limitations. The pre-MoU sample is small (211 incidents with "
     "valid SWH). The SEs on beta_3 (~0.20) imply we could only detect "
     "effects larger than ~0.40 (exp(0.40) = 1.49, a 49% gradient increase) "
     "at 5% significance. Moderate gradient changes would go undetected."),
    ("Pre-MoU gradient instability. The gradient varied substantially across "
     "pre-MoU years (Wald p = 0.083). If the baseline gradient was already "
     "noisy, detecting a shift at the MoU date is difficult."),
]
for pt in points:
    doc.add_paragraph(pt, style="List Bullet")

doc.add_heading("6.3 What This Does NOT Mean", level=2)
doc.add_paragraph(
    "A null weather gradient interaction does not mean the MoU did not make "
    "the route more dangerous. The MoU may have increased danger through "
    "channels that Model A cannot detect: more incidents overall (extensive "
    "margin), deaths in calm conditions (not captured by the gradient), "
    "or systematic under-detection of the worst outcomes."
)

# ── 7. Open Questions ──────────────────────────────────────────────
doc.add_heading("7. Open Questions and Next Steps", level=1)

doc.add_heading("7.1 Analyses Not Yet Performed", level=2)
points = [
    "Monthly-level analysis: did total monthly deaths increase post-MoU? "
    "This captures the extensive margin (more incidents) that Model A, "
    "by conditioning on incidents, cannot see.",
    "Crossings diagnostic: did the MoU change the volume of crossings in "
    "response to weather? (crossings as outcome, not conditioning variable)",
    "DML (double machine learning): partially linear model for beta_3 with "
    "flexible ML nuisance functions, to check whether the log-link functional "
    "form drives the null result.",
    "Detection sensitivity analysis: parameterized (pi, d-bar) grid to assess "
    "how much detection must have worsened to mask a positive beta_3.",
    "Re-run 02_build_event_data.R with the hmax grid index fix to recover "
    "max wave height data (currently 1,362/1,381 NAs due to a bug, now fixed).",
]
for pt in points:
    doc.add_paragraph(pt, style="List Bullet")

doc.add_heading("7.2 Thesis Direction Decision", level=2)
doc.add_paragraph(
    "The null weather gradient result is a genuine finding. The thesis can proceed "
    "on the following path:"
)

points = [
    "Report the Model A null as the per-incident severity result (the MoU did "
    "not amplify weather sensitivity).",
    "Add monthly-level analysis for the aggregate effect (total deaths, total "
    "incidents, controlling for crossings volume).",
    "Interpret the combination: if total deaths increased but per-incident "
    "weather sensitivity did not, the mechanism is volume (more incidents) "
    "and/or non-weather mortality (calm-sea deaths), not amplified weather risk.",
    "Support with descriptive evidence: geographic shift toward Libya, "
    "survivor ratio decline, weather overlap shift toward calmer conditions.",
]
for i, pt in enumerate(points, 1):
    doc.add_paragraph(f"{i}. {pt}")

doc.add_heading("7.3 Data Integrity Note", level=2)
doc.add_paragraph(
    "Four incidents (0.3% of sample) have 'Number Dead' > 'Total Dead and Missing' "
    "in the IOM data. This reflects IOM internal inconsistencies between separate "
    "fields, not a coding error. The analysis uses IOM's curated 'Total Dead and "
    "Missing' field (dead_missing) as the authoritative outcome. The largest "
    "discrepancy is the April 18, 2015 shipwreck (dead = 750, dead_missing = 472), "
    "reflecting different source estimates for the same event."
)

# ── 8. Technical Notes ─────────────────────────────────────────────
doc.add_heading("8. Technical Notes", level=1)

points = [
    "Estimation: fixest::fenegbin (R). Properly handles high-dimensional FE via "
    "demeaning, automatic singleton drops, heteroskedasticity-robust and "
    "cluster-robust SEs.",
    "Previous convergence issues with MASS::glm.nb (alternation limit reached "
    "with 188 grid FE levels) are resolved by switching to fixest.",
    "Grid FE: 1-degree grid (84 cells, 25 singletons dropped). "
    "Sensitivity tested at 0.25-degree (317 cells) and no grid FE.",
    "Weather extraction: ERA5 daily data at incident coordinates, nearest "
    "grid cell lookup. A bug in 02_build_event_data.R used 0.25-deg atmospheric "
    "grid indices for the 0.5-deg wave_hmax file, causing 98.6% NAs. "
    "Fixed (line 260: ft == 'wave' -> ft %in% c('wave', 'wave_hmax')), "
    "pending re-extraction.",
]
for pt in points:
    doc.add_paragraph(pt, style="List Bullet")

# ── Save ────────────────────────────────────────────────────────────
out_path = os.path.join(BASE_DIR, "documents", "model_a_progress_report.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
